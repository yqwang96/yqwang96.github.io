from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "files" / "WANGYinquan-CV-2026.docx"

# compact_reference_guide + named override: a4_two_page_cv
FONT = "Arial Unicode MS"
NAVY = "17324D"
TEAL = "0E7C7B"
INK = "1F2933"
MUTED = "5E6B78"
LIGHT = "E8F1F3"
RULE = "B8C8D1"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cellless_font(run, size=8.3, color=INK, bold=False, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def set_keep_with_next(paragraph, value=True):
    ppr = paragraph._p.get_or_add_pPr()
    node = ppr.find(qn("w:keepNext"))
    if value and node is None:
        ppr.append(OxmlElement("w:keepNext"))
    elif not value and node is not None:
        ppr.remove(node)


def set_keep_lines(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    if ppr.find(qn("w:keepLines")) is None:
        ppr.append(OxmlElement("w:keepLines"))


def set_bottom_border(paragraph, color=RULE, size=5, space=2):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def set_repeatable_bullet_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "360")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "300")
    ind.set(qn("w:hanging"), "150")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "8")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.extend([tabs, ind, spacing])
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    rpr.append(fonts)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, ppr, rpr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, numid])


def add_field(paragraph, field):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, text, end])
    set_cellless_font(run, size=7.5, color=MUTED)


def add_section_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(3.2)
    p.paragraph_format.space_after = Pt(1.4)
    set_keep_with_next(p)
    run = p.add_run(text)
    set_cellless_font(run, size=9.8, color=NAVY, bold=True)
    set_bottom_border(p)
    return p


def add_project_header(doc, title, role, date, tags=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1.4)
    p.paragraph_format.space_after = Pt(0.2)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(18.1), WD_TAB_ALIGNMENT.RIGHT)
    set_keep_with_next(p)
    set_keep_lines(p)
    r = p.add_run(title)
    set_cellless_font(r, size=8.9, color=NAVY, bold=True)
    r = p.add_run(f"｜{role}")
    set_cellless_font(r, size=7.7, color=MUTED, bold=True)
    r = p.add_run(f"\t{date}")
    set_cellless_font(r, size=7.7, color=MUTED, bold=True)
    if tags:
        q = doc.add_paragraph()
        q.paragraph_format.space_before = Pt(0)
        q.paragraph_format.space_after = Pt(0.2)
        set_keep_with_next(q)
        r = q.add_run(tags)
        set_cellless_font(r, size=7.1, color=TEAL, bold=True)


def add_bullet(doc, num_id, text, *, size=8.0, after=0.25):
    p = doc.add_paragraph()
    apply_num(p, num_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    set_keep_lines(p)
    r = p.add_run(text)
    set_cellless_font(r, size=size, color=INK)
    return p


def add_labeled_line(doc, label, text, *, size=7.95, after=0.35):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(label)
    set_cellless_font(r, size=size, color=NAVY, bold=True)
    r = p.add_run(text)
    set_cellless_font(r, size=size, color=INK)
    return p


def add_experience(doc, company, role, date, description):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0.8)
    p.paragraph_format.space_after = Pt(0.1)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(18.1), WD_TAB_ALIGNMENT.RIGHT)
    set_keep_with_next(p)
    r = p.add_run(company)
    set_cellless_font(r, size=8.45, color=NAVY, bold=True)
    r = p.add_run(f"｜{role}")
    set_cellless_font(r, size=7.65, color=MUTED, bold=True)
    r = p.add_run(f"\t{date}")
    set_cellless_font(r, size=7.65, color=MUTED, bold=True)
    q = doc.add_paragraph()
    q.paragraph_format.space_before = Pt(0)
    q.paragraph_format.space_after = Pt(0.45)
    q.paragraph_format.line_spacing = 1.0
    r = q.add_run(description)
    set_cellless_font(r, size=7.9, color=INK)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(8.3)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0.4)
    normal.paragraph_format.line_spacing = 1.0

    for style_name, size, color in [
        ("Heading 1", 9.8, NAVY),
        ("Heading 2", 9.1, NAVY),
        ("Heading 3", 8.6, TEAL),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(1.4)
        style.paragraph_format.keep_with_next = True


def configure_section(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.9)
    section.bottom_margin = Cm(0.9)
    section.left_margin = Cm(1.05)
    section.right_margin = Cm(1.05)
    section.header_distance = Cm(0.45)
    section.footer_distance = Cm(0.45)


def add_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("王印权  ·  Agentic AI / Reinforcement Learning / Industrial Intelligence  ·  ")
    set_cellless_font(r, size=6.6, color=MUTED)
    add_field(p, "PAGE")


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    configure_section(section)
    add_footer(section)
    num_id = set_repeatable_bullet_numbering(doc)

    # Compact memo_masthead variant; no decorative bottom border.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("王 印 权")
    set_cellless_font(r, size=19.5, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0.8)
    r = p.add_run("Agentic AI  ·  强化学习  ·  工业智能化与模型部署")
    set_cellless_font(r, size=9.2, color=TEAL, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.2)
    r = p.add_run("北京  |  +86-198-0120-8423  |  yinquanwang@qq.com  |  github.com/yqwang96  |  yqwang96.github.io")
    set_cellless_font(r, size=7.5, color=MUTED)

    add_section_heading(doc, "个人简介")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0.5)
    p.paragraph_format.line_spacing = 1.0
    text = (
        "系统科学博士，聚焦 Agentic AI、强化学习与工业智能化落地。具备企业级 Agent Runtime/Harness、"
        "Context Engineering、RAG/GraphRAG、Agent RLVR 及 CUDA/昇腾异构推理服务经验；完成供热控制、"
        "煤矿安全等项目从需求分析、数据治理、模型训练、服务部署、系统集成到生产效果复盘的端到端交付。"
        "博士阶段研究强化学习与运筹优化，发表 SCI 论文4篇。"
    )
    r = p.add_run(text)
    set_cellless_font(r, size=8.0, color=INK)

    add_section_heading(doc, "专业技能")
    add_labeled_line(doc, "Agent 系统：", "Agent Runtime/Harness、ReAct、CodeAct、Planner–Executor、Multi-Agent、Context Engineering、Tool Protocol、Skills、RAG/GraphRAG、Checkpoint、Human-in-the-loop")
    add_labeled_line(doc, "强化学习与后训练：", "DQN、PPO、Actor-Critic、多智能体强化学习、GRPO、RLVR、LoRA/PEFT、Verifier、Reward Shaping、Curriculum、Trajectory Evaluation")
    add_labeled_line(doc, "数据科学与运筹：", "LightGBM、XGBoost、时序预测、异常检测、PyTorch、PuLP/CBC、CPLEX/DOcplex、调度、匹配、策略优化与仿真评估")
    add_labeled_line(doc, "部署与工程：", "Python、SQL、FastAPI/Flask、React、PostgreSQL/MySQL、Docker、vLLM、Ascend-vLLM、CUDA、CANN/torch-npu、WebSocket/SSE、OPC UA")

    add_section_heading(doc, "工作经历")
    add_experience(
        doc,
        "中国联合网络通信有限公司北京人工智能科技中心",
        "数据科学家",
        "2025.08 — 至今",
        "负责 Agent 系统架构、LLM 应用、预测模型与强化学习研发，参与业务需求梳理、系统集成、生产部署及上线效果评估。",
    )
    add_experience(
        doc,
        "联通数科数据智能事业部",
        "数据科学家",
        "2024.07 — 2025.08",
        "参与工业互联网、时序预测、异常检测与智能体项目，完成算法原型、模型服务和业务应用的工程化落地。",
    )

    add_section_heading(doc, "核心项目")
    add_project_header(
        doc,
        "企业级数据科学 Agent 平台",
        "核心开发",
        "2025.06 — 至今",
        "Agent Runtime · Context Engineering · Tool/Skill · Docker Sandbox · vLLM/Ascend-vLLM",
    )
    add_bullet(doc, num_id, "面向团队数据分析与业务交付研发企业级 Agent 平台，支持自然语言驱动的数据探索、代码执行、机器学习建模、可视化与报告生成，支撑团队100+成员日常使用。")
    add_bullet(doc, num_id, "在不同交付场景中实践单主 Agent Harness 与 Orchestrator 驱动的多 Agent 架构，设计任务规划、工具调度、状态管理、执行反馈、错误恢复与终止判断机制。")
    add_bullet(doc, num_id, "构建分层 Context Engineering 体系，动态组织系统规则、工具 Schema、项目文件、执行输出、历史经验和压缩摘要，支撑长任务连续执行与上下文裁剪。")
    add_bullet(doc, num_id, "设计 Tool Protocol、Skill 与 Workspace 机制，接入 Python、SQL、检索、文件、模型训练和报告工具，形成多用户、多项目及文件/数据集/模型/Artifact 资产闭环。")
    add_bullet(doc, num_id, "实现后台任务、WebSocket/SSE 进度推送、断线续传、检查点恢复和持久化审批；通过 Docker 沙箱、路径约束、危险操作拦截和资源限制保障自动执行安全。")
    add_bullet(doc, num_id, "在 NVIDIA CUDA 和华为昇腾环境部署 vLLM/Ascend-vLLM 推理服务，以 OpenAI 兼容 API 接入 Agent，统一管理模型、Base URL、Context Window 与 Provider 配置。")

    add_project_header(
        doc,
        "供热预测控制与工业智能体",
        "核心开发",
        "2025.05 — 至今",
        "LightGBM · Qwen2.5 LoRA · Text-to-SQL · OPC UA · 安全控制闭环",
    )
    add_bullet(doc, num_id, "面向供热企业建设预测控制与工业智能体平台，打通 OPC UA、室温采集和气象 API，接入8,000+工业点位，覆盖130+换热站和200+二次网回路。")
    add_bullet(doc, num_id, "基于 LightGBM 构建站点级供水温度和循环泵频率小模型，融合历史工况、天气、室温、周期与相似工况特征；生产模型 R² 达0.98，室温控制偏差≤±0.5°C。")
    add_bullet(doc, num_id, "从144个站点、3个供暖季构建8,185条领域样本，基于 Qwen2.5-0.5B 与 LoRA 开展领域适应；温度 MAE 由3.43°C降至1.10°C、频率 MAE 由2.12Hz降至0.51Hz，结构化解析率由57%提升至100%。")
    add_bullet(doc, num_id, "构建供热问数 Agent，支持 Schema 自动同步、业务知识维护、Text-to-SQL、多轮工具调用、结果解释和图表生成；以只读权限、SQL 校验、行数限制和审计日志控制查询风险。")
    add_bullet(doc, num_id, "搭建异常诊断 Skill 框架，将异常记录与温度、压力、频率、能耗、天气、室温和历史控制指令关联，输出异常等级、可能根因和排查建议；知识规则与工具流程已完成，现场闭环验证持续迭代。")
    add_bullet(doc, num_id, "建立“预测—策略生成—安全校验—人工审批—指令下发—运行监控—偏差诊断”闭环，通过 dry-run、调整步长、设备模式检查和失败回滚约束控制风险；系统实现约7%节热、至少10%节电、投诉量下降5%。")

    page_break = doc.add_paragraph()
    page_break.paragraph_format.space_before = Pt(0)
    page_break.paragraph_format.space_after = Pt(0)
    page_break.add_run().add_break(WD_BREAK.PAGE)

    add_project_header(
        doc,
        "OR-Agent-R1：面向运筹优化的 Agent RLVR 训练系统",
        "独立研究",
        "2026.06 — 至今",
        "Qwen2.5-7B · LoRA/GRPO · RLVR · PuLP/CBC · Verifier · 6×A100",
    )
    add_bullet(doc, num_id, "构建运筹优化 Agent 的结构化执行环境，通过 XML Action Protocol 支持问题抽取、数学建模、PuLP 代码生成、CBC 求解、错误修正和最终回答。")
    add_bullet(doc, num_id, "将动作格式、状态转移、代码执行、Solver Status 与 Objective Match 拆分为分层可验证奖励，并采用 solve_format→solver→objective Curriculum 缓解工具型 Agent 的奖励稀疏。")
    add_bullet(doc, num_id, "基于 Qwen2.5-7B-Instruct、LoRA 与 GRPO 完成6×A100分布式训练，建立训练轨迹诊断、Checkpoint 选择、错误分类和逐题配对评估流程。")
    add_bullet(doc, num_id, "针对环境依赖、动作解析、代码兼容、Verifier 偏差与 Reward Hacking 建立 Canary 测试；完善 grouped-question 切分、Split Manifest 和跨来源重复审计，降低评估泄漏风险。")

    add_project_header(
        doc,
        "煤矿瓦斯异常征兆识别系统",
        "核心研发",
        "2025.02 — 2025.10",
        "多传感器时序 · Recurrence Plot/CWT · ResNet18 · ClickHouse · Flask · Docker",
    )
    add_bullet(doc, num_id, "面向煤矿安全生产构建多传感器异常检测系统，识别煤与瓦斯突出、冒顶片帮等6类典型异常征兆，完成需求分析、数据治理、模型研发、服务部署和现场效果评估。")
    add_bullet(doc, num_id, "将72小时多变量时序转换为递归图和 CWT 时频图，设计多路 ResNet18 双通道融合模型，刻画异常前兆的波动、极值和趋势演化；模型 AUC 达96.65%，召回率较单视图基线提升105%。")
    add_bullet(doc, num_id, "将模型封装为在线预警服务，对接 ClickHouse 实时数据流，通过 Flask API 与 Docker 完成生产部署、日志监控和业务系统集成；经阈值优化与现场反馈，误报率下降25%。")
    add_bullet(doc, num_id, "核心模型服务覆盖全省30+重点工作面，相关方案后续推广至33+矿、340+工作面，形成“实时接入—特征构建—在线推理—风险预警—现场反馈”的端到端闭环。")

    add_section_heading(doc, "强化学习与运筹优化研究")
    add_bullet(doc, num_id, "博士阶段围绕网约车订单派发、车辆调度和多平台协作开展强化学习与运筹优化研究，建立需求生成、订单匹配、车辆状态转移、策略训练和多算法对比的仿真评估体系。")
    add_bullet(doc, num_id, "设计融合长期价值估计与业务约束的订单派发、车辆重调度策略，实验中将订单响应率提升4.15%、平台收入提升7.30%。")
    add_bullet(doc, num_id, "研究路径由经典 DQN/PPO/Actor-Critic 与多智能体强化学习，延伸到工具型 Agent 的 GRPO/RLVR、可执行环境、Verifier 和可验证奖励。")

    add_section_heading(doc, "教育背景")
    add_experience(doc, "北京交通大学｜系统科学学院", "系统科学 · 硕博连读", "2018.09 — 2024.06", "研究方向：网约出行市场运营管理、强化学习与运筹优化；系统科学一流学科。")
    add_experience(doc, "日本广岛大学｜国际协力研究科", "国家公派访问博士生", "2022.05 — 2023.05", "研究方向：网约出行市场运营管理与交通运输规划。")
    add_experience(doc, "青岛理工大学｜机械与汽车工程学院", "交通工程 · 本科", "2014.09 — 2018.06", "")

    add_section_heading(doc, "代表论文")
    add_bullet(doc, num_id, "Reassignment Algorithm of the Ride-Sourcing Market Based on Reinforcement Learning，IEEE Transactions on Intelligent Transportation Systems，2023。", size=7.45, after=0.15)
    add_bullet(doc, num_id, "Promoting Collaborative Dispatching in the Ride-Sourcing Market with a Third-Party Integrator，IEEE Transactions on Intelligent Transportation Systems，2024。", size=7.45, after=0.15)
    add_bullet(doc, num_id, "Reinforcement Learning-Based Order-Dispatching Optimization in Ride-Sourcing Service，Computers & Industrial Engineering，2024。", size=7.45, after=0.15)
    add_bullet(doc, num_id, "Order Dispatching Optimization in Ride-Sourcing Market by Considering Cross Service Modes，Journal of Central South University，2023。", size=7.45, after=0.15)

    add_section_heading(doc, "荣誉奖励")
    add_bullet(doc, num_id, "2025年度奋斗者，中国联通数科。", size=7.45, after=0.1)
    add_bullet(doc, num_id, "第二届数字中国创新大赛全国二等奖（队伍负责人，排序1/3），2020。", size=7.45, after=0.1)
    add_bullet(doc, num_id, "中国研究生数学建模竞赛全国三等奖（队伍负责人），2019。", size=7.45, after=0.1)
    add_bullet(doc, num_id, "国家留学基金委公派联合培养博士研究生奖学金，2021。", size=7.45, after=0.1)

    doc.core_properties.title = "王印权简历 2026"
    doc.core_properties.subject = "Agentic AI、强化学习、工业智能化与模型部署"
    doc.core_properties.author = "王印权"
    doc.core_properties.keywords = "Agentic AI, Reinforcement Learning, Industrial Intelligence, Model Deployment"
    doc.core_properties.comments = ""
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
